import pandas as pd
from docx import Document
import os
import sys

def fill_template(excel_path, word_template_path, output_folder):
    xls = pd.ExcelFile(excel_path)
    first_sheet_df = pd.read_excel(xls, sheet_name=xls.sheet_names[0])

    for index, row in first_sheet_df.iterrows():
        doc = Document(word_template_path)

        for paragraph in doc.paragraphs:
            for key in first_sheet_df.columns:
                placeholder = f'{{{{{key}}}}}'
                value = row[key]

                if pd.isna(value):
                    for sheet_name in xls.sheet_names[1:]:
                        other_df = pd.read_excel(xls, sheet_name=sheet_name)
                        if index < len(other_df) and key in other_df.columns:
                            value = other_df.at[index, key]
                            if not pd.isna(value):
                                break

                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        for table in doc.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for key in first_sheet_df.columns:
                        placeholder = f'{{{{{key}}}}}'
                        value = row[key]

                        if pd.isna(value):
                            for sheet_name in xls.sheet_names[1:]:
                                other_df = pd.read_excel(xls, sheet_name=sheet_name)
                                if index < len(other_df) and key in other_df.columns:
                                    value = other_df.at[index, key]
                                    if not pd.isna(value):
                                        break

                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        output_filename = f'Agreement_Row_{index + 1}.docx'
        output_path = os.path.join(output_folder, output_filename)
        doc.save(output_path)
        print(f'Saved: {output_path}')

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python new_app.py <excel_path> <word_template_path> <output_folder>")
        sys.exit(1)

    excel_path = sys.argv[1]
    word_template_path = sys.argv[2]
    output_folder = sys.argv[3]

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    fill_template(excel_path, word_template_path, output_folder)